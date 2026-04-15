from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path("docs/ARBEITSANWEISUNG-WARENEINGANG.docx")

PRIMARY = RGBColor(0x12, 0x3B, 0x64)
MUTED = RGBColor(0x5A, 0x67, 0x76)
TEXT = RGBColor(0x1F, 0x29, 0x37)
BORDER = "D6E0EA"
FILL = "EEF4F8"
NOTICE_FILL = "F8FBFD"
NOTICE_BORDER = "9FB3C8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER, size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = TEXT


def add_paragraph(doc, text="", *, size=11, color=TEXT, bold=False, space_after=8, space_before=0, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(13)
    r.font.color.rgb = PRIMARY
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), BORDER)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_info_table(doc):
    table = doc.add_table(rows=2, cols=2)
    table.columns[0].width = Cm(3.8)
    table.columns[1].width = Cm(12.4)
    style_table(table)
    data = [
        ("Zielgruppe", "Mitarbeitende im Wareneingang und in der Fertigung"),
        ("Seite", 'Wareneingang in QIN Production Web'),
    ]
    for row, (left, right) in zip(table.rows, data):
        row.cells[0].width = Cm(3.8)
        row.cells[1].width = Cm(12.4)
        row.cells[0].paragraphs[0].clear()
        run = row.cells[0].paragraphs[0].add_run(left)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = PRIMARY
        set_cell_shading(row.cells[0], FILL)
        row.cells[1].text = right


def add_notice(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.columns[0].width = Cm(16.2)
    style_table(table)
    cell = table.cell(0, 0)
    set_cell_border(cell, NOTICE_BORDER)
    set_cell_shading(cell, NOTICE_FILL)
    cell.paragraphs[0].clear()
    r1 = cell.paragraphs[0].add_run("Wichtig: ")
    r1.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = PRIMARY
    r2 = cell.paragraphs[0].add_run(text)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = TEXT
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)


def add_data_table(doc, headings, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headings))
    style_table(table)
    for idx, width in enumerate(widths_cm):
        table.columns[idx].width = Cm(width)
    hdr = table.rows[0].cells
    for idx, heading in enumerate(headings):
        hdr[idx].width = Cm(widths_cm[idx])
        set_cell_shading(hdr[idx], FILL)
        hdr[idx].paragraphs[0].clear()
        run = hdr[idx].paragraphs[0].add_run(heading)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = PRIMARY
    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].width = Cm(widths_cm[idx])
            row[idx].text = value


def add_image_box(doc, title, note):
    table = doc.add_table(rows=1, cols=1)
    table.columns[0].width = Cm(16.2)
    style_table(table)
    cell = table.cell(0, 0)
    set_cell_border(cell, NOTICE_BORDER)
    set_cell_shading(cell, NOTICE_FILL)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    r1 = p.add_run(f"{title}: ")
    r1.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(10)
    r1.font.color.rgb = PRIMARY
    r2 = p.add_run(f"[Hier Screenshot einfügen] {note}")
    r2.font.name = "Calibri"
    r2.font.size = Pt(10)
    r2.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(0)


def set_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.clear()
    run = p.add_run("QIN Production Web | Arbeitsanweisung Wareneingang | v1.0")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    set_footer(section)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    add_paragraph(doc, "Arbeitsanweisung Wareneingang", size=19, color=PRIMARY, bold=True, space_after=2)
    add_paragraph(doc, "Fertigung | Version 1.0 | Stand: 15.04.2026", size=10.5, color=MUTED, space_after=8)

    add_info_table(doc)
    add_notice(doc, "Nur das Wichtigste: Was muss gemacht werden und worauf muss vor dem Speichern geachtet werden?")

    add_heading(doc, "Kurzablauf")
    add_data_table(
        doc,
        ["Schritt", "Was ist zu tun?"],
        [
            ("1", "Lieferant auswählen. Unten wird die Liste gefiltert."),
            ("2", "Passenden Eintrag über EBE-Nummer und Position auswählen."),
            ("3", "Pflichtfelder prüfen und ausfüllen."),
            ("4", "Vorhandene Chargen prüfen oder fehlende Chargen neu erfassen."),
            ("5", "Für jede Charge Etikett drucken und auf den Karton kleben."),
            ("6", "Erst dann speichern."),
        ],
        [3.0, 13.2],
    )
    add_image_box(doc, "Bild 1", "Lieferant auswählen und passenden Eintrag unten finden.")

    add_heading(doc, "Ablauf")
    add_data_table(
        doc,
        ["Schritt", "Was tun?"],
        [
            ("1", "Lieferant auswählen und richtigen Eintrag unten auswählen."),
            ("2", "Lieferschein, Zustand, Palettentausch und Dickenmessung prüfen."),
            ("3", "Wenn nötig Bemerkung eintragen."),
            ("4", "Wenn Chargen vorhanden sind: Charge und Menge mit dem Karton vergleichen."),
            ("5", "Wenn keine Chargen vorhanden sind: scannen oder vorsichtig manuell eintragen."),
            ("6", "Jede Charge markieren, Etikett drucken und auf den Karton kleben."),
            ("7", 'Am Ende auf "Änderungen Speichern" oder "Wareneingang Buchen" klicken.'),
        ],
        [3.0, 13.2],
    )
    add_image_box(doc, "Bild 2", "Eintrag laden, Pflichtfelder prüfen und Chargen kontrollieren.")

    add_heading(doc, "Vor dem Speichern prüfen")
    add_data_table(
        doc,
        ["Punkt", "Worauf achten?"],
        [
            ("Lieferant", "Muss ausgewählt sein."),
            ("Lieferschein", "Muss eingetragen sein."),
            ("Eintrag unten", "Der richtige Eintrag muss ausgewählt sein."),
            ("Dickenmessung", "Muss zwischen 0,23 mm und 1,2 mm liegen."),
            ("Bemerkung", "Ist Pflicht bei Palettentausch = Ja oder Zustand = Schlecht."),
            ("Chargen", "Stimmen Charge und Menge wirklich?"),
            ("Etiketten", "Für jede Charge Etikett drucken und auf den passenden Karton kleben."),
        ],
        [4.2, 12.0],
    )

    add_heading(doc, "Wichtig")
    add_data_table(
        doc,
        ["Situation", "Was tun?"],
        [
            ("Zustand = Schlecht", "Immer eine klare Bemerkung eintragen."),
            ("Palettentausch = Ja", "Immer eine Bemerkung eintragen."),
            ("Charge falsch", "Falsche Charge aus der Liste entfernen und neu erfassen."),
            ("Menge unklar", "Vor dem Speichern noch einmal prüfen."),
            ("Keine Charge vorhanden", "Charge scannen oder manuell eintragen. Nicht falsch eintippen."),
            ("Druck klappt nicht", "Erst prüfen, ob die richtige Charge markiert ist."),
        ],
        [4.8, 11.4],
    )
    add_image_box(doc, "Bild 3", "Chargen prüfen bzw. neu erfassen und Etikett drucken.")

    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    main()
